#!/usr/bin/env python3
"""Build Artie graduation report (.docx) from Project_Report_Artie.md — HU template style."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image as PilImage
except ImportError:
    PilImage = None

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / 'Project_Report_Artie.md'
OUT_PATH = ROOT / 'Artie_Graduation_Report.docx'

FIGURE_WIDTH = Inches(6.5)
TOC_MARKER = '<!-- TOC -->'

FIGURE_PLACEHOLDER_RE = re.compile(
    r'^\[FIGURE PLACEHOLDER[^\]]*\]|^\[PLACEHOLDER[^\]]*\]',
    re.IGNORECASE,
)

FRONT_MATTER_HEADINGS = frozenset(
    {
        'CERTIFICATE',
        'ABSTRACT',
        'Table of Contents',
        'LIST OF FIGURES',
        'LIST OF TABLES',
        'REFERENCES',
    }
)


def _embed_picture(run, image_path: Path, width=FIGURE_WIDTH) -> None:
    try:
        run.add_picture(str(image_path), width=width)
        return
    except Exception:
        if PilImage is None:
            raise
    with PilImage.open(image_path) as image:
        converted = image_path.with_suffix('.embed.png')
        image.convert('RGB').save(converted, format='PNG')
    run.add_picture(str(converted), width=width)


def _resolve_image(path_text: str) -> Path | None:
    cleaned = path_text.strip().split('?')[0]
    if cleaned.startswith('/'):
        candidate = Path(cleaned)
        if candidate.is_file():
            return candidate
    candidate = ROOT / cleaned
    if candidate.is_file():
        return candidate
    alt = ROOT / 'docs' / 'report_assets' / Path(cleaned).name
    return alt if alt.is_file() else None


def _add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def _shade_paragraph(paragraph, fill: str = 'FFF3CD') -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    p_pr.append(shd)


def _add_front_matter_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text.strip())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _add_toc_field(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-3" \h \z \u '

    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')

    placeholder = OxmlElement('w:t')
    placeholder.text = 'Right-click here and choose Update Field to refresh the Table of Contents.'

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)


def _add_figure_placeholder(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    message = (
        f'{text.strip()}\n'
        'Replace with your figure: Insert → Pictures. Center the image and add the caption below.'
    )
    run = paragraph.add_run(message)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    _shade_paragraph(paragraph, 'FFF3CD')


def _add_runs(paragraph, text: str) -> None:
    text = text.replace('$', '')
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    for r, row in enumerate(rows):
        for c in range(cols):
            value = row[c].strip() if c < len(row) else ''
            table.rows[r].cells[c].text = value


def _is_single_line_div(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith('<div') and stripped.endswith('</div>')


def _process_html_block(doc: Document, html_lines: list[str]) -> None:
    for html_line in html_lines:
        html_stripped = html_line.strip()
        if not html_stripped or html_stripped in {'<br>', '<br/>', '<br />'}:
            continue

        alt = ''
        src = ''
        md_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', html_stripped)
        html_match = re.search(r'<img[^>]+src=["\']([^"\']+)["\']', html_stripped)
        if md_match:
            alt, src = md_match.group(1), md_match.group(2)
        elif html_match:
            src = html_match.group(1)
            alt_match = re.search(r'alt=["\']([^"\']*)["\']', html_stripped)
            alt = alt_match.group(1) if alt_match else ''

        if src:
            resolved = _resolve_image(src)
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if resolved:
                run = paragraph.add_run()
                _embed_picture(run, resolved)
            else:
                _add_figure_placeholder(doc, f'[FIGURE PLACEHOLDER — missing file: {src}]')
            if alt:
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
            continue

        if html_stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(html_stripped[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
        elif html_stripped.startswith('### '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(p, html_stripped[4:].strip())
        elif html_stripped.startswith('<i>') and html_stripped.endswith('</i>'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(html_stripped.replace('<i>', '').replace('</i>', '')).italic = True
        elif not html_stripped.startswith('<'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(p, html_stripped)


def _add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_heading(text.strip(), level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def build_docx(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding='utf-8').splitlines()
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    index = 0
    in_code = False
    table_rows: list[list[str]] = []
    in_table = False
    in_html_block = False
    html_buffer: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped == TOC_MARKER:
            _add_toc_field(doc)
            index += 1
            continue

        if _is_single_line_div(line):
            if 'page-break' in stripped:
                _add_page_break(doc)
            index += 1
            continue

        if stripped.startswith('<div') and not _is_single_line_div(line):
            in_html_block = True
            html_buffer = []
            index += 1
            continue

        if in_html_block:
            if stripped.startswith('</div>'):
                _process_html_block(doc, html_buffer)
                html_buffer = []
                in_html_block = False
            else:
                html_buffer.append(line)
            index += 1
            continue

        if stripped.startswith('```'):
            in_code = not in_code
            index += 1
            continue
        if in_code:
            index += 1
            continue

        if not stripped:
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            index += 1
            continue

        if stripped.startswith('|'):
            cells = [cell.strip() for cell in stripped.strip('|').split('|')]
            if all(set(cell) <= {'-', ':', ' '} for cell in cells):
                index += 1
                continue
            table_rows.append(cells)
            in_table = True
            index += 1
            continue

        if in_table and table_rows:
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False

        if FIGURE_PLACEHOLDER_RE.match(stripped):
            _add_figure_placeholder(doc, stripped)
            index += 1
            continue

        img_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            alt, src = img_match.groups()
            resolved = _resolve_image(src)
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if resolved:
                run = paragraph.add_run()
                _embed_picture(run, resolved)
            else:
                _add_figure_placeholder(doc, f'[FIGURE PLACEHOLDER — {alt or src}]')
            if alt:
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            index += 1
            continue

        if stripped.startswith('# Chapter'):
            _add_page_break(doc)
            _add_heading(doc, stripped[2:].strip(), level=1)
        elif stripped.startswith('# '):
            heading_text = stripped[2:].strip()
            if heading_text in FRONT_MATTER_HEADINGS:
                _add_front_matter_heading(doc, heading_text)
            else:
                _add_heading(doc, heading_text, level=1)
        elif stripped.startswith('## '):
            _add_heading(doc, stripped[3:].strip(), level=2)
        elif stripped.startswith('### '):
            _add_heading(doc, stripped[4:].strip(), level=3)
        elif stripped.startswith('#### '):
            _add_heading(doc, stripped[5:].strip(), level=4)
        elif stripped.startswith('- '):
            paragraph = doc.add_paragraph(style='List Bullet')
            _add_runs(paragraph, stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            paragraph = doc.add_paragraph(style='List Number')
            _add_runs(paragraph, re.sub(r'^\d+\.\s', '', stripped))
        elif stripped.startswith('*Table') or stripped.startswith('*Figure'):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(stripped.lstrip('*').rstrip('*')).italic = True
        else:
            paragraph = doc.add_paragraph()
            _add_runs(paragraph, stripped)

        index += 1

    if table_rows:
        _add_table(doc, table_rows)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    words = len(md_path.read_text(encoding='utf-8').split())
    print(f'Wrote {out_path}')
    print(f'Markdown words: ~{words} (target ~9000+ with figures for 50 pages)')
    print(f'Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')


def main() -> int:
    md = MD_PATH if len(sys.argv) < 2 else Path(sys.argv[1])
    out = OUT_PATH if len(sys.argv) < 3 else Path(sys.argv[2])
    if not md.is_file():
        print(f'Missing markdown: {md}', file=sys.stderr)
        return 1
    build_docx(md, out)
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
